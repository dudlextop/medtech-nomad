from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

sys.path.append(str(Path(__file__).resolve().parents[1]))
from pipeline_common import as_str, build_raw_record, clean_price, ensure_dirs, load_sources, write_json, RAW_DIR, REPORTS_DIR


def parse_docx_source(source: dict) -> tuple[list[dict], dict]:
    path = Path(source["source_file"])
    records: list[dict] = []
    warnings: list[str] = []
    try:
        doc = Document(path)
        if not doc.tables:
            return [], {"source_file": str(path), "status": "warning", "records": 0, "warnings": ["No tables found"]}
        table = doc.tables[0]
        for idx, row in enumerate(table.rows[2:], start=3):
            cells = [as_str(cell.text).replace("\n", " ") for cell in row.cells]
            if len(cells) < 3:
                continue
            service_code, raw_name, price_raw = cells[0], cells[1], cells[2]
            if not raw_name or raw_name.lower().startswith("раздел"):
                continue
            price = clean_price(price_raw)
            rec = build_raw_record(
                source,
                service_code=service_code,
                raw_service_name=raw_name,
                unit=None,
                price_kzt=price,
                price_type="base",
                raw_row_json={"row_number": idx, "cells": cells},
                parser_confidence=0.98 if price else 0.25,
                row_index=idx,
            )
            if rec:
                records.append(rec)
            else:
                warnings.append(f"Skipped row {idx}: could not parse price/name")
    except Exception as exc:
        return [], {"source_file": str(path), "status": "failed", "records": 0, "warnings": [str(exc)]}

    report = {
        "source_file": str(path),
        "clinic_id": source.get("clinic_id"),
        "parser": "docx_table",
        "status": "success" if records else "warning",
        "records": len(records),
        "warnings": warnings[:100],
    }
    return records, report


def parse_all_docx() -> dict:
    ensure_dirs()
    all_records: list[dict] = []
    reports: list[dict] = []
    for source in load_sources():
        if source.get("parser_type") != "docx_table":
            continue
        records, report = parse_docx_source(source)
        all_records.extend(records)
        reports.append(report)
    write_json(RAW_DIR / "docx_price_records.json", all_records)
    write_json(REPORTS_DIR / "docx_import_report.json", {"files": reports, "records": len(all_records)})
    return {"records": len(all_records), "files": reports}


if __name__ == "__main__":
    print(parse_all_docx())
